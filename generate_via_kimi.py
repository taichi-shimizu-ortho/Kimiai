import json
import os
from openai import OpenAI
from dotenv import load_dotenv
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image

# 1. APIと環境変数の読み込み
load_dotenv()
api_key = os.getenv("KIMI_API_KEY")
base_url = os.getenv("KIMI_BASE_URL", "https://api.moonshot.ai/v1")
model = os.getenv("KIMI_MODEL", "kimi-k3")

if not api_key or api_key == "your_api_key_here":
    print("エラー: .env ファイルに正しい KIMI_API_KEY を設定してください。")
    exit(1)

client = OpenAI(api_key=api_key, base_url=base_url)

# 2. JSONデータの読み込み
json_path = r"C:\Users\a2189\Dropbox\obsidian\40_manuscript\2026ORS\domon.json"
with open(json_path, "r", encoding="utf-8") as f:
    prompt_data = f.read()

import re

# 3. Kimi API を呼び出して原稿を生成
messages = [
    {
        "role": "system", 
        "content": "あなたは優秀な医療・研究系のライターです。ユーザーから提供されたJSON形式のアウトラインをもとに、学会参加報告書の原稿を作成してください。です・ます調の自然な日本語で、適度に段落分けを行い、フォーマルなトーンで記述してください。出力は本文のみとしてください（タイトルや著者名などは不要です）。\n\n【重要】JSONデータ内に写真(photos)の指定がある場合、原稿内のその出来事を説明している段落の直後に、必ず `[PHOTO: ファイル名]` という形式のマーカーを挿入してください。（例: [PHOTO: photo_01_combined.jpg]）"
    },
    {"role": "user", "content": prompt_data}
]

print(f"KIMI API ({model}) に原稿作成を依頼しています...")
try:
    response = client.chat.completions.create(
        model=model,
        messages=messages,
        temperature=1.0,
    )
    manuscript_text = response.choices[0].message.content
    print("原稿の生成が完了しました！\n")
except Exception as e:
    print(f"APIエラーが発生しました: {e}")
    exit(1)

# 4. Wordドキュメントの生成
OUTPUT_PATH = r"C:\Users\a2189\Dropbox\obsidian\40_manuscript\2026ORS\ORS2026_参加報告_KIMI生成版.docx"
# resize_photos.py で作成したフォルダを参照
RESIZED_DIR = os.path.join(os.path.dirname(__file__), "resized_photos")

print("Wordファイルを作成中...")
document = Document()

# B5サイズ設定
section = document.sections[0]
section.page_width = Inches(7.17)
section.page_height = Inches(10.12)
section.left_margin, section.right_margin = Inches(0.8), Inches(0.8)
section.top_margin, section.bottom_margin = Inches(0.8), Inches(0.8)

# タイトル・著者名
title = document.add_heading('Orthopaedic Research Society 2026 Annual Meeting 参加報告', 1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
author = document.add_paragraph('大学院3年目 清水')
author.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# JSONから写真のキャプション辞書を作成しておく
data = json.loads(prompt_data)
photo_dict = {}
for day in data.get("schedule", []):
    for photo in day.get("photos", []):
        if "filename" in photo:
            photo_dict[photo["filename"]] = photo.get("caption", "")

# 生成された本文を追加 (マーカーを検出して写真を挿入)
for paragraph in manuscript_text.split('\n'):
    p_text = paragraph.strip()
    if not p_text:
        continue
        
    # [PHOTO: filename] のマーカーを探す
    match = re.search(r'\[PHOTO:\s*(.+?)\]', p_text)
    if match:
        filename = match.group(1).strip()
        
        # マーカーより前のテキストを追加
        text_before = p_text[:match.start()].strip()
        if text_before:
            document.add_paragraph(text_before)
            
        # 写真を挿入
        resized_path = os.path.join(RESIZED_DIR, filename)
        if os.path.exists(resized_path):
            document.add_picture(resized_path, width=Inches(3.5))
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            caption = photo_dict.get(filename, "")
            caption_p = document.add_paragraph(caption)
            caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_p.runs[0].font.size = Pt(9)
        else:
            document.add_paragraph(f"（※画像ファイルが見つかりません: {filename}）")
            
        # マーカーより後のテキストを追加
        text_after = p_text[match.end():].strip()
        if text_after:
            document.add_paragraph(text_after)
    else:
        document.add_paragraph(p_text)

document.save(OUTPUT_PATH)
print(f"ドキュメントの保存が完了しました: {OUTPUT_PATH}")
