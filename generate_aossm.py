import json
import os
import re
from openai import OpenAI
from dotenv import load_dotenv
from docx import Document
from docx.shared import Inches, Pt, Mm, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image

# Directories
target_dir = r"C:\Users\a2189\Dropbox\obsidian\00_daily\2026AOSSM"
gdrive_dir = r"G:\マイドライブ\2026AOSSM"

# Photos to resize
photo_filenames = ["1783606421573.jpg", "1783663381730.jpg", "P_20260709_082226.jpg"]
resized_photos = []

print("Resizing photos...")
for f in photo_filenames:
    src = os.path.join(gdrive_dir, f)
    dst = os.path.join(target_dir, f"resized_{f}")
    if os.path.exists(src):
        with Image.open(src) as img:
            # Correct image orientation if needed based on EXIF
            try:
                from PIL import ImageOps
                img = ImageOps.exif_transpose(img)
            except Exception:
                pass
            img.thumbnail((800, 800))
            img.save(dst, "JPEG", quality=85)
            resized_photos.append(dst)
            print(f"Resized {f} to {dst}")
    else:
        print(f"Photo not found: {src}")

# Load Kimi API
load_dotenv()
api_key = os.getenv("KIMI_API_KEY")
base_url = os.getenv("KIMI_BASE_URL", "https://api.moonshot.ai/v1")
model = os.getenv("KIMI_MODEL", "kimi-k3")

client = OpenAI(api_key=api_key, base_url=base_url)

# Read JSON
json_path = os.path.join(target_dir, "rep.json.md")
with open(json_path, "r", encoding="utf-8") as f:
    prompt_data = f.read()

# Call Kimi
messages = [
    {
        "role": "system", 
        "content": "あなたは優秀な医療・研究系のライターです。提供されたJSON形式のアウトラインをもとに、AOSSM参加報告書の原稿を作成してください。WordのB5サイズ1〜2ページ（約1200〜1600字程度）に収まるように、です・ます調の自然な日本語で、適度に段落分けを行い、フォーマルなトーンで記述してください。出力は本文のみとしてください（タイトルや著者名などは不要です）。\n\nまた、写真が3枚あります（1: 初日の木崎先生・内田先生とのディナー、2: UCSFの会、3: 学会会場）。原稿内の適切な箇所に `[PHOTO: 1]`, `[PHOTO: 2]`, `[PHOTO: 3]` という形式でマーカーを挿入してください。"
    },
    {"role": "user", "content": prompt_data}
]

print(f"Calling Kimi API ({model})...")
response = client.chat.completions.create(
    model=model,
    messages=messages,
    temperature=1.0,
)
manuscript_text = response.choices[0].message.content
print("API returned content.")

# Create Document
doc_path = os.path.join(target_dir, "AOSSM2026_参加報告.docx")
document = Document()

# B5 size setup (JIS B5 is 182mm x 257mm)
section = document.sections[0]
section.page_width = Mm(182)
section.page_height = Mm(257)
section.left_margin, section.right_margin = Mm(20), Mm(20)
section.top_margin, section.bottom_margin = Mm(20), Mm(20)

title = document.add_heading('AOSSM 2026 Annual Meeting 参加報告', 1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
author = document.add_paragraph('大学院3年目 清水')
author.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Map indices to actual resized paths based on visual guess or generic assignment
# 1783606421573, 1783663381730, P_20260709_082226
# I will just assign them as 1, 2, 3 in order of the array. The user can swap them in Word if they are mismatched.
photo_paths = {
    "1": resized_photos[0] if len(resized_photos) > 0 else None,
    "2": resized_photos[1] if len(resized_photos) > 1 else None,
    "3": resized_photos[2] if len(resized_photos) > 2 else None,
}
photo_captions = {
    "1": "初日の木崎先生・内田先生とのディナー",
    "2": "UCSFの会",
    "3": "学会会場"
}

# Parse and insert
for paragraph in manuscript_text.split('\n'):
    p_text = paragraph.strip()
    if not p_text:
        continue
        
    # Match [PHOTO: 1] or similar
    match = re.search(r'\[PHOTO:\s*(\d+)\]', p_text)
    if match:
        idx = match.group(1)
        text_before = p_text[:match.start()].strip()
        if text_before:
            document.add_paragraph(text_before)
            
        p_path = photo_paths.get(idx)
        if p_path and os.path.exists(p_path):
            try:
                document.add_picture(p_path, width=Cm(8))
                last_p = document.paragraphs[-1]
                last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                cap = photo_captions.get(idx, "")
                cap_p = document.add_paragraph(cap)
                cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_p.runs[0].font.size = Pt(9)
            except Exception as e:
                document.add_paragraph(f"（※画像 {idx} の挿入に失敗しました: {e}）")
        else:
            document.add_paragraph(f"（※画像 {idx} 挿入位置: ファイルが見つかりません）")
            
        text_after = p_text[match.end():].strip()
        if text_after:
            document.add_paragraph(text_after)
    else:
        document.add_paragraph(p_text)

document.save(doc_path)
print(f"Document saved to {doc_path}")
