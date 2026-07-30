from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os

doc_path = r"C:\Users\a2189\Dropbox\obsidian\00_daily\2026AOSSM\AOSSM2026_参加報告.docx"
img2_path = r"C:\Users\a2189\Dropbox\obsidian\00_daily\2026AOSSM\resized_1783663381730.jpg"

doc = Document(doc_path)
for i, p in enumerate(doc.paragraphs):
    if "[PHOTO: 2]" in p.text:
        # Remove the marker from the text
        p.text = p.text.replace("[PHOTO: 2]", "").strip()
        
        # Insert the image at the next paragraph position
        # Since python-docx doesn't easily let you insert a picture *between* paragraphs,
        # we can just clear a paragraph and add a run with the picture, or add a new paragraph.
        # However, to insert at a specific index, we can insert a paragraph before the next one.
        
        if i + 1 < len(doc.paragraphs):
            new_p = doc.paragraphs[i+1].insert_paragraph_before()
        else:
            new_p = doc.add_paragraph()
            
        run = new_p.add_run()
        run.add_picture(img2_path, width=Cm(8))
        new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add caption
        if i + 1 < len(doc.paragraphs):
            cap_p = doc.paragraphs[i+1].insert_paragraph_before("UCSFの会")
        else:
            cap_p = doc.add_paragraph("UCSFの会")
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_p.runs[0].font.size = Pt(9)
        break

doc.save(doc_path)
print("Fixed docx.")
