from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import os

# --- Configurations ---
OUTPUT_PATH = r"C:\Users\a2189\Dropbox\obsidian\40_manuscript\2026ORS\ORS2026_参加報告_清水.docx"
PHOTO_DIR = r"G:\マイドライブ\2026ORS写真"
TEMP_DIR = r"C:\Users\a2189\uv-envs\Kimiai\workspace\temp_photos"

# Selected photos
photos = [
    {"file": "photo_07_poster.jpg", "caption": "ポスター発表の様子"},
    {"file": "photo_05_asahikawa.jpg", "caption": "旭川の先生方との写真"}
]

# Create temp dir for resized photos
os.makedirs(TEMP_DIR, exist_ok=True)

def resize_image(input_path, output_path, max_size=800):
    with Image.open(input_path) as img:
        img.thumbnail((max_size, max_size))
        img.save(output_path, "JPEG", quality=85)

document = Document()

# Page Setup for B5 (JIS B5 is approx 7.17 x 10.12 inches)
section = document.sections[0]
section.page_width = Inches(7.17)
section.page_height = Inches(10.12)
section.left_margin = Inches(0.8)
section.right_margin = Inches(0.8)
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.8)

# Title
title = document.add_heading('Orthopaedic Research Society 2026 Annual Meeting 参加報告', 1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Author
author = document.add_paragraph('大学院3年目 清水')
author.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Text content
text = """2026年3月27日から3月31日にかけて、アメリカ合衆国ノースカロライナ州シャーロットで開催されたOrthopaedic Research Society 2026 Annual Meeting (ORS 2026) に参加いたしました。ORSはアメリカ国内に約3200人の会員を有する整形外科基礎分野の学会組織であり、今回は約2000題のポスター発表と約300題の口演発表が行われる非常に規模の大きな学会でした。例年多くの日本人研究者が参加していることでも知られています。

同門会からは、これまで2019年に藤谷先生、山中先生、川崎先生が、2022年には山中先生が参加され、参加報告を伺っておりました。私自身としては、自身の研究成果を国際学会で発表し、整形外科基礎研究の最新の知見に触れることを目的に参加いたしました。

学会開催地であるシャーロットは、地方都市らしい落ち着いた雰囲気があり、初めての訪問地でしたが非常に過ごしやすい街でした。気候としては、3月でも日中は暖かいものの、夜間は冷え込むような環境でした。

到着日および学会前日は、長時間の移動の疲れを癒やしつつ、翌日からの学会に向けた準備を行いました。
学会1日目には、さっそく会場へと足を運びました。夜には内田先生と合流し、また偶然にもPatrick Quinnさんとお会いする機会があり、大変有意義な時間を過ごしました。
2日目の昼には、「Mentor Mentee Luncheon」という参加型のプログラムに参加しました。ORSではこうしたプログラムも充実しており、各テーブルにメンターが1人つき、参加者が自由に研究内容やキャリアについて話し合う形式で、非常に刺激的な経験となりました。
3日目には、Flash poster sessionでの発表を行いました。夜には旭川の先生方とも交流する機会に恵まれました。
そして4日目には、自身のポスター発表として「股関節におけるリラキシンレセプター発現について」の発表を行いました。

ポスター発表では、自身の研究内容について英語で説明を行いました。国際学会という大きな舞台で自分の研究を発信できたことは、大変貴重な経験となりました。

今回の学会参加を通じて、整形外科基礎研究の非常に幅広いテーマに触れることができ、自分自身の研究の位置づけを改めて客観的に再確認する良い機会となりました。今回得られた知見と経験を、今後の研究活動や発表に大いに活かしていきたいと考えております。

最後になりますが、今回の学会参加にあたり、日頃から温かいご指導と多大なるご支援をいただいた先生方に深く感謝申し上げます。"""

for p in text.split('\n\n'):
    document.add_paragraph(p)

# Photos
for p_info in photos:
    orig_path = os.path.join(PHOTO_DIR, p_info["file"])
    if os.path.exists(orig_path):
        resized_path = os.path.join(TEMP_DIR, p_info["file"])
        resize_image(orig_path, resized_path)
        
        # Add to document
        document.add_picture(resized_path, width=Inches(3.5))
        
        # Add caption
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        caption_p = document.add_paragraph(p_info["caption"])
        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_p.runs[0].font.size = Pt(9)

document.save(OUTPUT_PATH)
print(f"Document saved successfully to {OUTPUT_PATH}")
